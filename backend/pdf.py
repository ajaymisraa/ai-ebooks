import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO
import subprocess

def download_image(image_url):
    response = requests.get(image_url)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        raise Exception("Error downloading image")

def convert_to_pdf(word_filename):
    pdf_filename = word_filename.replace('.docx', '.pdf')
    command = f"docx2pdf {word_filename} {pdf_filename}"
    subprocess.run(command, shell=True)
    return pdf_filename

def toc(json_input):
    # Parse the JSON string into a dictionary if it's a string
    if isinstance(json_input, str):
        json_input = json.loads(json_input)

    toc = {}
    for chapter, content in json_input.items():
        if chapter.startswith('Chapter'):
            subchapters = [subchapter for subchapter in content.keys()]
            toc[chapter] = subchapters
    return toc

def create_footer_with_page_numbers(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    field_code = 'PAGE   \\* MERGEFORMAT'
    docx_field = OxmlElement('w:fldSimple')
    docx_field.set(qn('w:instr'), field_code)
    run._r.append(docx_field)

def create_ebook(cover_image_path_or_url, book_contents, publisher):
    if isinstance(book_contents, str):
        book_contents = json.loads(book_contents)

    table_of_contents = toc(book_contents)
    title = book_contents['title']

    doc = Document()
    # Set the default font of the document
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'  # San Francisco is not available, using Arial instead
    style.font.size = Pt(11)

    # Add title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(36)  # Set the title font size
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add cover image
    if cover_image_path_or_url.startswith('http://') or cover_image_path_or_url.startswith('https://'):
        image_stream = download_image(cover_image_path_or_url)
        image = doc.add_picture(image_stream, width=Inches(4))
    else:
        image = doc.add_picture(cover_image_path_or_url, width=Inches(4))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ensure that we add the publisher after the image
    publisher_paragraph = doc.add_paragraph(f"{publisher}")
    publisher_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Move the publisher name to the bottom of the page
    section = doc.sections[0]
    sectPr = section._sectPr
    if sectPr is not None:
        footer = sectPr.xpath('./w:footerReference')
        if footer:
            footer[0].getparent().remove(footer[0])
    footer_section = doc.sections[0]
    footer_paragraph = footer_section.footer.paragraphs[0]
    footer_paragraph.text = f"{publisher} | "
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page numbers to the footer
    for section in doc.sections:
        create_footer_with_page_numbers(section)

    doc.add_page_break()

    # Add Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for chapter_title, subchapters in table_of_contents.items():
        doc.add_paragraph(chapter_title, style='List Number')
        for subchapter in subchapters:
            doc.add_paragraph(subchapter, style='List Bullet')

    doc.add_page_break()

    # Add Book Contents
    for chapter_title, sections in book_contents.items():
        if chapter_title != 'title':  # Exclude the title from the chapter processing
            doc.add_heading(chapter_title, level=1)
            for section_title, paragraphs in sections.items():
                doc.add_heading(section_title, level=2)
                for paragraph in paragraphs:
                    doc.add_paragraph(paragraph['paragraph'])


    word_filename = 'ebook.docx'
    doc.save(word_filename)
    print("Word document saved.")
    pdf_filename = convert_to_pdf(word_filename)
    print("PDF saved.")
    return pdf_filename


#with open('test.json', 'r') as file:
#   json_input = json.load(file)

# Generate Table of Contents
# table_of_contents = toc(json_input)
